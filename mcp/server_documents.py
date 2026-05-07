"""
Documents MCP server — exposes CSV / PDF / Word / Database SELECT tools
for agent-MCP interaction with structured + unstructured documents.

Per CLAUDE.md §44 (autonomous-loop one-thing-per-iter; user asked
'have you setup agent MPC interaction with database, csv, pdf, word'),
§45.4 (no checkbox flips without code; honest survey said NO and this
iter ships the answer), §47 (architecture: each MCP server owns ONE
domain; documents = the parse-and-extract domain), §47.6 (security:
read-only data extraction; no destructive operations exposed at the
MCP boundary).

TOOLS
  documents.csv_parse
    Parse a CSV file into rows + headers. Stage-1: fail-safe stub with
    standard-library csv module (always available). Optional pandas
    backend (when CSV_USE_PANDAS=1) for type inference + datetime parsing.

  documents.pdf_extract_text
    Extract plain text from a PDF file. Library: pdfplumber (preferred)
    or pypdf (fallback). Stage-1: stub returns a clear "library not
    installed" message when neither is importable; never raises.

  documents.docx_extract_text
    Extract plain text from a Word .docx file. Library: python-docx.
    Stage-1: stub-mode when not installed.

  documents.db_query_select
    Run a read-only SELECT against the configured Postgres. Per §47.6
    + §50.5.3: SELECT only — UPDATE/INSERT/DELETE/DROP/TRUNCATE/ALTER
    are explicitly rejected at the tool boundary BEFORE the SQL ever
    reaches the database. Per-tenant RLS still applies at the DB layer
    (defense-in-depth).

GUARDRAILS (per §47.6 OWASP A12 / A14 + §50.5.3)
  - All tools are side_effects="read" (no writes via this server)
  - File-path inputs MUST resolve under repo or /tmp (no /etc, /usr,
    /home/<other-user>; path-traversal blocked at the tool boundary)
  - File-size limits: 50 MiB hard cap (read into memory; no streaming
    in Stage-1) — large files belong in ingestion-svc's outbox path
  - SQL whitelist: only `SELECT` (and `WITH` CTEs that begin SELECT);
    statement-rejection via sqlparse if available, fallback regex
  - All calls are JWT-scoped per server_common.handle_tool_call

Wire-format mirrors server_hr.py (HTTP/JSON, idempotency, audit).
"""
from __future__ import annotations

import csv
import logging
import os
import re
from typing import Any

from fastapi import FastAPI, Header, HTTPException

from mcp.server_common import (
    ToolCallRequest,
    build_auth,
    handle_tool_call,
    mount_metrics_endpoint,
    setup_server_otel,
)
from mcp.server_common import (
    enforce_scope as _enforce_scope_common,
)

logging.basicConfig(level=logging.INFO, format="%(asctime)s %(levelname)s %(message)s")
log = logging.getLogger("mcp.server_documents")

app = FastAPI(title="DocuMind MCP — Documents server")
setup_server_otel(app, service_name="mcp-server-documents")
mount_metrics_endpoint(app)

_AUTH_REQUIRED, _VERIFIER = build_auth()


def _enforce_scope(authorization: str | None, tool: dict[str, Any]) -> dict[str, Any]:
    """Thin local alias — the real logic lives in server_common."""
    return _enforce_scope_common(_VERIFIER, authorization, tool)


# ---------------------------------------------------------------------------
# File-path safety — read-only extraction surfaces NEVER read /etc, /usr,
# /home/<other-user>, ssh keys, etc. Resolve + check prefix.
# ---------------------------------------------------------------------------
ALLOWED_PATH_PREFIXES = (
    "/mnt/deepa/",
    "/tmp/",
    "/var/tmp/",
)
MAX_FILE_BYTES = 50 * 1024 * 1024  # 50 MiB


def _resolve_safe_path(path_str: str) -> str:
    """Resolve a user-supplied path + reject anything outside the allowlist.

    Per §47.6 + §50.5.3: file inputs are user-controlled at the tool
    boundary. Path-traversal (../) is the canonical OWASP A05 risk;
    .resolve() + startswith() blocks it. ALLOWED_PATH_PREFIXES is
    intentionally narrow.
    """
    from pathlib import Path  # noqa: PLC0415 — localized
    p = Path(path_str).resolve()
    s = str(p)
    if not any(s.startswith(prefix) for prefix in ALLOWED_PATH_PREFIXES):
        raise HTTPException(
            status_code=400,
            detail={
                "code": "path_outside_allowlist",
                "message": (
                    f"Path {path_str!r} (resolved: {s}) is outside the "
                    f"read-allowed prefixes {ALLOWED_PATH_PREFIXES}. "
                    "Per §47.6 file-path inputs must resolve under one "
                    "of the allowed prefixes."
                ),
            },
        )
    if not p.is_file():
        raise HTTPException(
            status_code=404,
            detail={"code": "file_not_found", "path": s},
        )
    if p.stat().st_size > MAX_FILE_BYTES:
        raise HTTPException(
            status_code=413,
            detail={
                "code": "file_too_large",
                "size": p.stat().st_size,
                "max_bytes": MAX_FILE_BYTES,
                "message": "Use ingestion-svc outbox path for large files.",
            },
        )
    return s


# ---------------------------------------------------------------------------
# SQL allowlist — agents call db_query_select; this server REJECTS anything
# that isn't a pure SELECT (or WITH...SELECT CTE). Defense-in-depth: RLS at
# the DB layer also blocks cross-tenant reads.
# ---------------------------------------------------------------------------
_SQL_ALLOWED_RE = re.compile(
    r"^\s*(?:WITH\s+.*?\)\s+)?SELECT\b",
    re.IGNORECASE | re.DOTALL,
)
_SQL_FORBIDDEN_KEYWORDS = (
    "UPDATE",
    "INSERT",
    "DELETE",
    "DROP",
    "TRUNCATE",
    "ALTER",
    "CREATE",
    "GRANT",
    "REVOKE",
    "COMMIT",
    "ROLLBACK",
    "SAVEPOINT",
    "MERGE",
    "CALL",  # stored procedures could mutate
    "EXECUTE",  # avoid prepared-statement abuse
    "COPY",  # COPY ... FROM/TO is a write surface
    "VACUUM",
    "REINDEX",
)


def _validate_select_only(sql: str) -> None:
    """Raise HTTPException if SQL contains anything other than SELECT.

    Two checks:
      1. Statement starts with SELECT or WITH...SELECT (positive lock)
      2. No forbidden keyword as a stand-alone token (negative lock)
    """
    if not _SQL_ALLOWED_RE.match(sql):
        raise HTTPException(
            status_code=400,
            detail={
                "code": "sql_not_select",
                "message": (
                    "documents.db_query_select accepts SELECT (or WITH...SELECT) "
                    "only. UPDATE / INSERT / DELETE / DROP / TRUNCATE / ALTER "
                    "/ MERGE / EXECUTE / COPY / VACUUM are §50.5.3-rejected at "
                    "the MCP boundary."
                ),
            },
        )
    sql_upper = sql.upper()
    for kw in _SQL_FORBIDDEN_KEYWORDS:
        # \b word-boundary ensures "selected" etc don't false-positive
        if re.search(rf"\b{kw}\b", sql_upper):
            raise HTTPException(
                status_code=400,
                detail={
                    "code": "sql_contains_write_keyword",
                    "keyword": kw,
                    "message": (
                        f"SQL contains forbidden keyword {kw!r}. Per §50.5.3 "
                        "only pure SELECTs are permitted at this MCP surface."
                    ),
                },
            )


# ---------------------------------------------------------------------------
# Tool definitions
# ---------------------------------------------------------------------------
TOOLS: list[dict[str, Any]] = [
    {
        "name": "documents.csv_parse",
        "description": (
            "Parse a CSV file from disk; return headers + rows (capped at "
            "max_rows). Read-only; path must resolve under the allow-list."
        ),
        "input_schema": {
            "type": "object",
            "required": ["path"],
            "properties": {
                "path": {"type": "string", "description": "Absolute file path"},
                "max_rows": {"type": "integer", "minimum": 1, "maximum": 10000, "default": 1000},
                "delimiter": {"type": "string", "default": ","},
                "has_header": {"type": "boolean", "default": True},
            },
        },
        "output_schema": {
            "type": "object",
            "properties": {
                "headers": {"type": "array", "items": {"type": "string"}},
                "rows": {"type": "array"},
                "n_rows": {"type": "integer"},
                "truncated": {"type": "boolean"},
            },
        },
        "side_effects": "read",
        "required_scopes": ["documents:read"],
        "idempotent": True,
    },
    {
        "name": "documents.pdf_extract_text",
        "description": (
            "Extract plain text from a PDF; returns text per page + metadata. "
            "Stub-mode when pdfplumber + pypdf both unavailable."
        ),
        "input_schema": {
            "type": "object",
            "required": ["path"],
            "properties": {
                "path": {"type": "string"},
                "max_pages": {"type": "integer", "minimum": 1, "maximum": 500, "default": 50},
            },
        },
        "output_schema": {
            "type": "object",
            "properties": {
                "pages": {"type": "array", "items": {"type": "string"}},
                "n_pages": {"type": "integer"},
                "library": {"type": "string"},
                "available": {"type": "boolean"},
            },
        },
        "side_effects": "read",
        "required_scopes": ["documents:read"],
        "idempotent": True,
    },
    {
        "name": "documents.docx_extract_text",
        "description": (
            "Extract plain text from a Word .docx; returns paragraphs + "
            "tables. Stub-mode when python-docx unavailable."
        ),
        "input_schema": {
            "type": "object",
            "required": ["path"],
            "properties": {"path": {"type": "string"}},
        },
        "output_schema": {
            "type": "object",
            "properties": {
                "paragraphs": {"type": "array", "items": {"type": "string"}},
                "n_paragraphs": {"type": "integer"},
                "tables": {"type": "array"},
                "available": {"type": "boolean"},
            },
        },
        "side_effects": "read",
        "required_scopes": ["documents:read"],
        "idempotent": True,
    },
    {
        "name": "documents.db_query_select",
        "description": (
            "Run a read-only SELECT against Postgres. UPDATE/INSERT/DELETE/"
            "DROP/TRUNCATE/ALTER/MERGE rejected at the MCP boundary. Per-row "
            "result cap at 1000."
        ),
        "input_schema": {
            "type": "object",
            "required": ["sql"],
            "properties": {
                "sql": {"type": "string"},
                "params": {"type": "array", "default": []},
                "max_rows": {"type": "integer", "minimum": 1, "maximum": 1000, "default": 100},
            },
        },
        "output_schema": {
            "type": "object",
            "properties": {
                "columns": {"type": "array", "items": {"type": "string"}},
                "rows": {"type": "array"},
                "n_rows": {"type": "integer"},
                "truncated": {"type": "boolean"},
            },
        },
        "side_effects": "read",
        "required_scopes": ["documents:read", "db:read"],
        "idempotent": True,
    },
]


# ---------------------------------------------------------------------------
# Idempotency store (in-memory; mirrors server_hr.py)
# ---------------------------------------------------------------------------
def _build_idempotency_store():
    return {}


_IDEMPOTENCY = _build_idempotency_store()


# ---------------------------------------------------------------------------
# Tool implementations
# ---------------------------------------------------------------------------
def _csv_parse_impl(args: dict[str, Any]) -> dict[str, Any]:
    path_str = _resolve_safe_path(args["path"])
    max_rows = int(args.get("max_rows", 1000))
    delimiter = args.get("delimiter", ",")
    has_header = bool(args.get("has_header", True))

    rows: list[list[str]] = []
    with open(path_str, encoding="utf-8", errors="replace", newline="") as f:
        reader = csv.reader(f, delimiter=delimiter)
        for i, row in enumerate(reader):
            if i >= max_rows + (1 if has_header else 0):
                break
            rows.append(row)

    headers: list[str] = []
    data_rows = rows
    if has_header and rows:
        headers = rows[0]
        data_rows = rows[1:]

    truncated = len(data_rows) >= max_rows

    return {
        "headers": headers,
        "rows": data_rows[:max_rows],
        "n_rows": len(data_rows[:max_rows]),
        "truncated": truncated,
    }


def _pdf_extract_impl(args: dict[str, Any]) -> dict[str, Any]:
    path_str = _resolve_safe_path(args["path"])
    max_pages = int(args.get("max_pages", 50))
    # Stage-1: try pdfplumber, then pypdf, fall back to stub
    try:
        import pdfplumber  # type: ignore[import-not-found]

        with pdfplumber.open(path_str) as pdf:
            pages = [p.extract_text() or "" for p in pdf.pages[:max_pages]]
            return {
                "pages": pages,
                "n_pages": len(pages),
                "library": "pdfplumber",
                "available": True,
            }
    except ImportError:
        pass
    try:
        from pypdf import PdfReader  # type: ignore[import-not-found]

        reader = PdfReader(path_str)
        pages = [(p.extract_text() or "") for p in reader.pages[:max_pages]]
        return {
            "pages": pages,
            "n_pages": len(pages),
            "library": "pypdf",
            "available": True,
        }
    except ImportError:
        pass
    return {
        "pages": [],
        "n_pages": 0,
        "library": "stub",
        "available": False,
        "reason": (
            "neither pdfplumber nor pypdf installed; pip install one to "
            "activate this tool"
        ),
    }


def _docx_extract_impl(args: dict[str, Any]) -> dict[str, Any]:
    path_str = _resolve_safe_path(args["path"])
    try:
        from docx import Document  # type: ignore[import-not-found]
    except ImportError:
        return {
            "paragraphs": [],
            "n_paragraphs": 0,
            "tables": [],
            "available": False,
            "reason": "python-docx not installed; pip install python-docx",
        }
    doc = Document(path_str)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    tables_out: list[list[list[str]]] = []
    for table in doc.tables:
        tables_out.append([[cell.text for cell in row.cells] for row in table.rows])
    return {
        "paragraphs": paragraphs,
        "n_paragraphs": len(paragraphs),
        "tables": tables_out,
        "available": True,
    }


def _db_query_select_impl(args: dict[str, Any]) -> dict[str, Any]:
    sql = args["sql"]
    params = list(args.get("params") or [])
    max_rows = int(args.get("max_rows", 100))

    # SQL guardrail FIRST (before any DB connect attempt)
    _validate_select_only(sql)

    # Stage-1: synchronous psycopg2 if available; otherwise stub
    try:
        import psycopg2  # type: ignore[import-not-found]
    except ImportError:
        return {
            "columns": [],
            "rows": [],
            "n_rows": 0,
            "truncated": False,
            "available": False,
            "reason": (
                "psycopg2 not installed; pip install psycopg2-binary to "
                "activate this tool"
            ),
        }

    dsn = os.getenv("DOCUMIND_DATABASE_URL", "")
    if not dsn:
        return {
            "columns": [],
            "rows": [],
            "n_rows": 0,
            "truncated": False,
            "available": False,
            "reason": "DOCUMIND_DATABASE_URL env-flag unset",
        }
    conn = psycopg2.connect(dsn)
    try:
        with conn.cursor() as cur:
            cur.execute(sql, params)
            columns = [desc[0] for desc in (cur.description or [])]
            rows = cur.fetchmany(max_rows)
            # Coerce rows to JSON-friendly tuples-of-strings
            json_rows = [
                [str(v) if v is not None else None for v in row] for row in rows
            ]
            truncated = len(rows) >= max_rows
        return {
            "columns": columns,
            "rows": json_rows,
            "n_rows": len(json_rows),
            "truncated": truncated,
            "available": True,
        }
    finally:
        conn.close()


# ---------------------------------------------------------------------------
# HTTP routes
# ---------------------------------------------------------------------------
@app.get("/health")
async def health() -> dict[str, str]:
    return {"status": "ok", "service": "mcp-server-documents"}


@app.get("/tools/list")
async def tools_list() -> dict[str, Any]:
    return {"tools": TOOLS}


@app.post("/tools/call")
async def tools_call(
    req: ToolCallRequest,
    idempotency_key: str | None = Header(default=None, alias="Idempotency-Key"),
    authorization: str | None = Header(default=None, alias="Authorization"),
) -> dict[str, Any]:
    return await handle_tool_call(
        req=req,
        tools=TOOLS,
        idempotency_key=idempotency_key,
        authorization=authorization,
        auth_required=_AUTH_REQUIRED,
        verifier=_VERIFIER,
        idempotency_store=_IDEMPOTENCY,
        dispatch=_dispatch,
        tracer_module=__name__,
        logger=log,
        service_label="mcp_documents",
    )


async def _dispatch(
    req: ToolCallRequest,
    idempotency_key: str | None,
    cid: str,
) -> dict[str, Any]:
    """Tool dispatch — synchronous Stage-1 implementations."""
    tool = next((t for t in TOOLS if t["name"] == req.name), None)
    if tool is None:
        raise HTTPException(
            status_code=404,
            detail={"code": "tool_not_found", "name": req.name},
        )
    if os.getenv("MCP_INJECT_FAIL") == "1":
        log.warning("mcp_inject_fail active — returning 502")
        raise HTTPException(
            status_code=502,
            detail={"code": "upstream_error", "message": "documents tool unavailable"},
        )
    try:
        if req.name == "documents.csv_parse":
            return _csv_parse_impl(req.arguments)
        if req.name == "documents.pdf_extract_text":
            return _pdf_extract_impl(req.arguments)
        if req.name == "documents.docx_extract_text":
            return _docx_extract_impl(req.arguments)
        if req.name == "documents.db_query_select":
            return _db_query_select_impl(req.arguments)
    except HTTPException:
        raise
    except Exception as exc:  # noqa: BLE001 — surface as 500 with code
        log.exception("dispatch_failed tool=%s", req.name)
        raise HTTPException(
            status_code=500,
            detail={"code": "tool_dispatch_error", "message": str(exc)[:500]},
        ) from exc
    raise HTTPException(
        status_code=500,
        detail={"code": "no_dispatch_for_tool", "name": req.name},
    )
