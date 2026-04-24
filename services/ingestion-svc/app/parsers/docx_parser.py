"""DOCX parser built on :mod:`python-docx`."""
from __future__ import annotations

import io
import logging

from docx import Document  # type: ignore[import-untyped]

from .base import DocumentParser, ParsedDocument, ParsedPage

log = logging.getLogger(__name__)


class DocxParser(DocumentParser):
    extensions = (".docx",)

    def parse(self, data: bytes, *, filename: str) -> ParsedDocument:
        warnings: list[str] = []
        try:
            doc = Document(io.BytesIO(data))
        except Exception as exc:
            log.warning("docx_parse_failed filename=%s err=%s", filename, exc)
            warnings.append(f"parse_failed: {exc}")
            return ParsedDocument(title=filename, pages=[], metadata={"parse_warnings": warnings})

        # DOCX has no native "pages". We split on explicit section breaks
        # (Heading 1 starts a new page-equivalent). Fallback: one big page.
        pages: list[list[str]] = [[]]
        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue
            if para.style.name.startswith("Heading 1") and pages[-1]:
                pages.append([])
            pages[-1].append(text)

        result_pages = [
            ParsedPage(page_number=i, text="\n".join(lines))
            for i, lines in enumerate(pages, start=1)
            if lines
        ]
        title = doc.core_properties.title or filename
        return ParsedDocument(
            title=str(title),
            pages=result_pages,
            metadata={
                "source_filename": filename,
                "section_count": len(result_pages),
                "parse_warnings": warnings,
            },
        )
